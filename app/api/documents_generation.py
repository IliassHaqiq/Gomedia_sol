"""
Document Generation API Endpoints

This module provides REST API endpoints for generating:
- A&E Specifications Word documents
- Datasheet ZIP archives

Supports both Excel file upload and manual part number input.
"""
import os
import uuid
import tempfile
import zipfile
import logging
import requests
from io import BytesIO
from datetime import datetime
from typing import List, Optional, Dict
import traceback

import pandas as pd
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
from sqlalchemy.orm import Session

from app.db.session import get_db
from app.models.product import Product, ProductDescription, ProductFile

router = APIRouter()

# Configuration du logging
logger = logging.getLogger(__name__)


def is_valid_file(file_path: str) -> bool:
    """Check if file exists (local path or R2 URL)."""
    if not file_path:
        return False

    if file_path.startswith('http'):
        try:
            response = requests.head(file_path, timeout=5, allow_redirects=True)
            return response.status_code == 200
        except Exception as e:
            logger.warning(f"URL check failed for {file_path}: {e}")
            return False
    else:
        return os.path.exists(file_path)


def extract_product_data_from_excel(excel_path: str) -> Dict[str, dict]:
    """
    Extract product data from Excel file.
    Expected columns: 'reference', 'N prix', 'designation'
    Returns: dict mapping reference -> {N prix, designation}
    """
    try:
        df = pd.read_excel(excel_path)
        df.columns = df.columns.str.strip()

        required_cols = {"reference", "N prix", "designation"}
        missing_cols = required_cols - set(df.columns)
        if missing_cols:
            raise HTTPException(
                status_code=400,
                detail=f"Excel file must contain columns: {', '.join(required_cols)}. Missing: {', '.join(missing_cols)}. Found: {', '.join(df.columns)}"
            )

        excel_data = {}
        for _, row in df.iterrows():
            ref = str(row["reference"]).strip()
            if not ref or pd.isna(ref):
                continue

            n_prix = str(row["N prix"]).strip() if pd.notna(row["N prix"]) else ""
            designation = str(row["designation"]).strip() if pd.notna(row["designation"]) else ""

            excel_data[ref] = {
                "N prix": n_prix,
                "designation": designation
            }

        return excel_data

    except Exception as e:
        if isinstance(e, HTTPException):
            raise
        raise HTTPException(status_code=400, detail=f"Invalid Excel file: {str(e)}")


def extract_part_numbers_simple(excel_path: str) -> List[str]:
    """Extract just part numbers from Excel (for manual endpoints)."""
    try:
        df = pd.read_excel(excel_path)
        df.columns = df.columns.str.strip()

        if "Part Number" in df.columns:
            parts = df["Part Number"].dropna().astype(str).str.strip().tolist()
        elif "reference" in df.columns:
            parts = df["reference"].dropna().astype(str).str.strip().tolist()
        else:
            raise HTTPException(
                status_code=400,
                detail="Excel file must contain 'Part Number' or 'reference' column"
            )

        cleaned_parts = [p for p in parts if p and p.lower() != "nan"]
        return cleaned_parts

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid Excel file: {str(e)}")


def fetch_products_from_db(db: Session, parts: List[str]) -> List[dict]:
    """Fetch product data for given part numbers."""
    if not parts:
        return []

    products = db.query(Product).filter(
        Product.ref_produit.in_(parts)
    ).all()

    result = []
    for product in products:
        description = db.query(ProductDescription).filter(
            ProductDescription.product_id == product.id
        ).first()

        result.append({
            "id": product.id,
            "ref_produit": product.ref_produit,
            "designation": product.designation,
            "descriptif_fr": description.descriptif_fr if description else None
        })

    return result


def fetch_products_and_files_from_db(db: Session, parts: List[str]) -> List[dict]:
    """Fetch product data along with associated file paths."""
    if not parts:
        return []

    products = db.query(Product).filter(
        Product.ref_produit.in_(parts)
    ).all()

    result = []
    for product in products:
        description = db.query(ProductDescription).filter(
            ProductDescription.product_id == product.id
        ).first()

        files = db.query(ProductFile).filter(
            ProductFile.product_id == product.id
        ).all()

        for file in files:
            result.append({
                "id": product.id,
                "ref_produit": product.ref_produit,
                "designation": product.designation,
                "descriptif_fr": description.descriptif_fr if description else None,
                "file_path": file.file_path
            })

    return result


def generate_ae_specs_docx(project_name: str, products: List[dict], excel_data: Dict[str, dict] = None) -> bytes:
    """
    Generate a Word document with A&E specifications.
    Now includes N prix from Excel if provided.
    """
    doc = Document()
    doc.add_heading("Spécifications A&E", 0)

    for product in products:
        ref_produit = product.get("ref_produit", "")
        designation_db = product.get("designation") or ""
        descriptif = product.get("descriptif_fr") or "(aucun descriptif disponible)"

        # Récupérer le N prix depuis Excel si disponible
        n_prix = ""
        if excel_data and ref_produit in excel_data:
            n_prix = excel_data[ref_produit].get("N prix", "")

        # Titre : Designation (Reference) - N prix
        title_parts = []
        if designation_db:
            title_parts.append(designation_db)
        if ref_produit:
            title_parts.append(f"({ref_produit})")
        if n_prix:
            title_parts.append(f"- N° prix: {n_prix}")

        title = " ".join(title_parts) if title_parts else "Produit"
        doc.add_heading(title, level=1)

        # Ajouter les détails
        if ref_produit:
            doc.add_paragraph(f"Référence: {ref_produit}")
        if n_prix:
            doc.add_paragraph(f"N° de prix: {n_prix}")

        doc.add_paragraph("\nDescription:")
        doc.add_paragraph(descriptif)

        # Espace entre produits
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_zip_from_files(project_name: str, file_urls: List[str], arcnames: List[str]) -> bytes:
    """Generate a ZIP archive from R2 URLs (downloads files on the fly)."""
    buffer = BytesIO()

    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        for i, url in enumerate(file_urls):
            if not url:
                logger.warning(f"Empty URL at index {i}, skipping")
                continue

            try:
                logger.debug(f"Downloading: {url}")
                response = requests.get(url, timeout=30)
                response.raise_for_status()

                zf.writestr(arcnames[i], response.content)
                logger.debug(f"Added to ZIP: {arcnames[i]}")
            except requests.RequestException as e:
                logger.warning(f"Failed to download {url}: {e}")
            except Exception as e:
                logger.warning(f"Error processing {url}: {e}")

    buffer.seek(0)
    return buffer.getvalue()


def clean_filename(s: str) -> str:
    """Remove invalid characters for filenames."""
    invalid_chars = '<>:"/\\|?*'
    for ch in invalid_chars:
        s = s.replace(ch, '_')
    return s.strip()


# Pydantic schemas
class PartsRequest(BaseModel):
    """Schema for manual part number request"""
    project_name: str
    parts: List[str]


@router.post("/ae-specs")
async def api_generate_ae_specs(
    excel: UploadFile = File(...),
    project_name: str = Form(...),
    db: Session = Depends(get_db)
):
    """
    Generate A&E Specs Word document from an Excel file.
    Excel must contain: reference, N prix, designation
    """
    logger.info(f"=== Génération A&E Specs : project_name='{project_name}' ===")

    if not project_name or not project_name.strip():
        raise HTTPException(status_code=400, detail="Project name is required")

    temp_dir = tempfile.mkdtemp()
    excel_path = os.path.join(temp_dir, f"{uuid.uuid4()}.xlsx")
    logger.debug(f"Dossier temporaire: {temp_dir}")

    try:
        # Sauvegarde du fichier Excel
        with open(excel_path, "wb") as f:
            content = await excel.read()
            f.write(content)
        logger.info(f"Fichier Excel sauvegardé: {excel_path} ({len(content)} bytes)")

        # Extraire les données de l'Excel
        excel_data = extract_product_data_from_excel(excel_path)
        references = list(excel_data.keys())
        logger.info(f"Références extraites ({len(references)}): {references[:10]}...")

        if not references:
            raise HTTPException(status_code=400, detail="No valid references found in Excel file")

        # Récupérer les produits depuis la DB
        products = fetch_products_from_db(db, references)
        logger.info(f"Produits trouvés en BD: {len(products)}")

        if not products:
            raise HTTPException(status_code=404, detail="No products found in database for the given references")

        # Générer le document Word avec les données Excel + DB
        logger.info("Génération du document Word...")
        docx_bytes = generate_ae_specs_docx(project_name.strip(), products, excel_data)
        logger.info(f"Document généré: {len(docx_bytes)} bytes")

        return StreamingResponse(
            BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{project_name.strip()}_A&E_Specs.docx"'
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ERREUR lors de la génération A&E Specs: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")
    finally:
        try:
            if os.path.exists(excel_path):
                os.unlink(excel_path)
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
            logger.debug("Nettoyage des fichiers temporaires OK")
        except Exception as e:
            logger.warning(f"Erreur lors du nettoyage: {e}")


@router.post("/datasheet-zip")
async def api_generate_datasheet_zip(
    excel: UploadFile = File(...),
    project_name: str = Form(...),
    db: Session = Depends(get_db)
):
    """
    Generate a ZIP file containing datasheet PDFs for products in the Excel file.
    Excel must contain: reference, N prix, designation
    """
    logger.info(f"=== Génération ZIP : project_name='{project_name}' ===")

    if not project_name or not project_name.strip():
        raise HTTPException(status_code=400, detail="Project name is required")

    temp_dir = tempfile.mkdtemp()
    excel_path = os.path.join(temp_dir, f"{uuid.uuid4()}.xlsx")
    logger.debug(f"Dossier temporaire: {temp_dir}")

    try:
        # Sauvegarde du fichier Excel
        with open(excel_path, "wb") as f:
            content = await excel.read()
            f.write(content)
        logger.info(f"Fichier Excel sauvegardé: {excel_path} ({len(content)} bytes")

        # Extraire les données de l'Excel
        excel_data = extract_product_data_from_excel(excel_path)
        references = list(excel_data.keys())
        logger.info(f"Références extraites ({len(references)}): {references[:10]}...")

        if not references:
            raise HTTPException(status_code=400, detail="No valid references found in Excel file")

        # Récupérer les produits et fichiers depuis la DB
        products_with_files = fetch_products_and_files_from_db(db, references)
        logger.info(f"Produits trouvés: {len(products_with_files)}")

        # Construire les listes de fichiers URLs et noms personnalisés
        file_urls = []
        arcnames = []

        for p in products_with_files:
            file_url = p.get("file_path")
            ref = p.get("ref_produit")

            if file_url and is_valid_file(file_url):
                file_urls.append(file_url)

                # Créer un nom personnalisé avec N prix, désignation, référence
                excel_row = excel_data.get(ref)
                if excel_row:
                    n_prix = excel_row.get("N prix", "")
                    designation_val = excel_row.get("designation", "")
                else:
                    n_prix = ""
                    designation_val = ""

                # Récupérer l'extension du fichier depuis l'URL
                _, ext = os.path.splitext(file_url)
                if not ext:
                    ext = ".pdf"

                clean_ref = clean_filename(ref)
                clean_n_prix = clean_filename(n_prix) if n_prix else ""
                clean_designation = clean_filename(designation_val) if designation_val else ""

                # Format: Nprix_Designation_reference.extension
                if clean_n_prix and clean_designation:
                    arcname = f"{clean_n_prix}_{clean_designation}_{clean_ref}{ext}"
                elif clean_designation:
                    arcname = f"{clean_designation}_{clean_ref}{ext}"
                else:
                    arcname = f"{clean_ref}{ext}"

                arcnames.append(arcname)
            else:
                logger.warning(f"Fichier inaccessible: {file_url}")

        logger.info(f"Fichiers PDF valides: {len(file_urls)}")
        if file_urls:
            logger.debug(f"Noms personnalisés: {arcnames[:3]}")

        if not file_urls:
            raise HTTPException(status_code=404, detail="No datasheet files found for the selected products")

        # Génération du ZIP
        logger.info("Génération du ZIP en cours...")
        zip_bytes = generate_zip_from_files(project_name.strip(), file_urls, arcnames)
        logger.info(f"ZIP généré: {len(zip_bytes)} bytes")

        return StreamingResponse(
            BytesIO(zip_bytes),
            media_type="application/zip",
            headers={
                "Content-Disposition": f'attachment; filename="{project_name.strip()}.zip"'
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ERREUR lors de la génération du ZIP: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error generating ZIP: {str(e)}")
    finally:
        try:
            if os.path.exists(excel_path):
                os.unlink(excel_path)
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
            logger.debug("Nettoyage des fichiers temporaires OK")
        except Exception as e:
            logger.warning(f"Erreur lors du nettoyage: {e}")


@router.post("/ae-specs/manual")
async def api_generate_ae_specs_manual(
    request: PartsRequest,
    db: Session = Depends(get_db)
):
    """
    Generate A&E Specs Word document from a list of part numbers (JSON body).
    Note: N prix not available in manual mode (only part numbers).
    """
    project_name = request.project_name
    part_numbers = request.parts

    logger.info(f"=== Génération A&E Specs (manuel) : project_name='{project_name}' ===")

    if not project_name or not project_name.strip():
        raise HTTPException(status_code=400, detail="Project name is required")

    if not part_numbers:
        raise HTTPException(status_code=400, detail="No part numbers provided")

    logger.info(f"Numéros de pièce fournis ({len(part_numbers)}): {part_numbers[:10]}...")

    try:
        products = fetch_products_from_db(db, part_numbers)
        logger.info(f"Produits trouvés: {len(products)}")

        if not products:
            raise HTTPException(status_code=404, detail="No products found for the given part numbers")

        # En mode manuel, pas de données Excel donc pas de N prix dans le Word
        logger.info("Génération du document Word...")
        docx_bytes = generate_ae_specs_docx(project_name.strip(), products, excel_data=None)

        return StreamingResponse(
            BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{project_name.strip()}_A&E_Specs.docx"'
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ERREUR lors de la génération A&E Specs (manuel): {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")


@router.post("/datasheet-zip/manual")
async def api_generate_datasheet_zip_manual(
    request: PartsRequest,
    db: Session = Depends(get_db)
):
    """
    Generate a ZIP file containing datasheet PDFs from a list of part numbers (JSON body).
    """
    project_name = request.project_name
    part_numbers = request.parts

    logger.info(f"=== Génération ZIP (manuel) : project_name='{project_name}' ===")

    if not project_name or not project_name.strip():
        raise HTTPException(status_code=400, detail="Project name is required")

    if not part_numbers:
        raise HTTPException(status_code=400, detail="No part numbers provided")

    logger.info(f"Numéros de pièce fournis ({len(part_numbers)}): {part_numbers[:10]}...")

    try:
        products_with_files = fetch_products_and_files_from_db(db, part_numbers)
        logger.info(f"Produits trouvés: {len(products_with_files)}")

        file_urls = []
        arcnames = []

        for p in products_with_files:
            file_url = p.get("file_path")
            ref = p.get("ref_produit")

            if file_url and is_valid_file(file_url):
                file_urls.append(file_url)

                # En mode manuel, pas de N prix ni désignation depuis Excel
                # Utiliser le nom de fichier de l'URL
                filename = os.path.basename(file_url)
                if filename:
                    arcnames.append(filename)
                else:
                    arcnames.append(f"{ref}.pdf")
            else:
                logger.warning(f"Fichier inaccessible: {file_url}")

        logger.info(f"Fichiers PDF valides: {len(file_urls)}")

        if not file_urls:
            raise HTTPException(status_code=404, detail="No datasheet files found for the selected products")

        logger.info("Génération du ZIP en cours...")
        zip_bytes = generate_zip_from_files(project_name.strip(), file_urls, arcnames)
        logger.info(f"ZIP généré: {len(zip_bytes)} bytes")

        return StreamingResponse(
            BytesIO(zip_bytes),
            media_type="application/zip",
            headers={
                "Content-Disposition": f'attachment; filename="{project_name.strip()}.zip"'
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ERREUR lors de la génération du ZIP (manuel): {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error generating ZIP: {str(e)}")
